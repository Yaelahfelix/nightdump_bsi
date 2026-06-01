"""
Generator Proposal Night Dump
Jalankan: python3 proposal_night_dump.py
Output: proposal_night_dump.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_paragraph_format(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          space_before=0, space_after=6,
                          line_spacing=1.5, first_indent=1.25):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_indent:
        pf.first_line_indent = Cm(first_indent)

def add_text(doc, text, size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True,
             space_before=0, space_after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size, bold, italic)
    set_paragraph_format(para, align, space_before, space_after,
                         first_indent=1.25 if indent else 0)
    return para

def add_heading_1(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=14, bold=True)
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.CENTER,
                         space_before=12, space_after=6,
                         first_indent=0)
    return para

def add_heading_2(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=12, bold=True)
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=6, space_after=3,
                         first_indent=0)
    return para

def add_heading_3(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=12, bold=True, italic=True)
    set_paragraph_format(para, WD_ALIGN_PARAGRAPH.LEFT,
                         space_before=3, space_after=3,
                         first_indent=0)
    return para

def add_bullet(doc, text, indent_cm=1.5):
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    set_font(run, size=12)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.5
    return para

def add_numbered(doc, text):
    para = doc.add_paragraph(style='List Number')
    run = para.add_run(text)
    set_font(run, size=12)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.5
    return para

def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        set_font(p.add_run(' '), size=6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

def add_table_row(table, row_idx, data, bold_first=False):
    row = table.rows[row_idx]
    for i, cell_text in enumerate(data):
        cell = row.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(cell_text)
        set_font(run, size=11, bold=(i == 0 and bold_first))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 0 else WD_ALIGN_PARAGRAPH.LEFT
    return row

# ─────────────────────────────────────────────────────────────────────────────
# MAIN GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def create_proposal():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin    = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin   = Cm(4)
        section.right_margin  = Cm(3)

    # ═══════════════════════════════════════════════════════════════════════════
    # HALAMAN JUDUL
    # ═══════════════════════════════════════════════════════════════════════════
    for _ in range(3):
        add_spacer(doc)

    p = doc.add_paragraph()
    r = p.add_run("PROPOSAL PROYEK AKHIR\nMATERI KULIAH MOBILE PROGRAMMING")
    set_font(r, 14, bold=True)
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    add_spacer(doc, 2)

    p = doc.add_paragraph()
    r = p.add_run("Night Dump")
    set_font(r, 24, bold=True)
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

    p = doc.add_paragraph()
    r = p.add_run("Aplikasi Manajemen Pikiran Malam\nBerbasis Kecerdasan Buatan")
    set_font(r, 14, italic=True)
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)
    add_spacer(doc, 3)

    for line in [
        "Diajukan untuk Memenuhi Tugas Mata Kuliah Mobile Programming",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_font(r, 12)
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

    add_spacer(doc, 2)

    p = doc.add_paragraph()
    r = p.add_run("Disusun oleh:")
    set_font(r, 12, bold=True)
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

    for line in [
        "Ferdinand Felix",
        "felix@mrksolution.id",
    ]:
        p = doc.add_paragraph()
        r = p.add_run(line)
        set_font(r, 12)
        set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

    add_spacer(doc, 4)

    p = doc.add_paragraph()
    r = p.add_run("PROGRAM STUDI TEKNOLOGI INFORMASI\nUNIVERSITAS BSI\n" + str(datetime.datetime.now().year))
    set_font(r, 12, bold=True)
    set_paragraph_format(p, WD_ALIGN_PARAGRAPH.CENTER, first_indent=0)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BAB I PENDAHULUAN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "BAB I\nPENDAHULUAN")

    add_heading_2(doc, "1.1 Latar Belakang")
    add_text(doc, (
        "Di era modern yang serba cepat, manusia seringkali dihadapkan pada "
        "kondisi pikiran yang tidak bisa beristirahat menjelang tidur. Fenomena "
        "ini dikenal sebagai overthinking atau racing mind, yaitu kondisi di mana "
        "seseorang terus-menerus memikirkan berbagai hal seperti tugas yang belum "
        "selesai, rencana esok hari, kekhawatiran, maupun ide-ide baru yang muncul "
        "tiba-tiba. Akibatnya, kualitas tidur menurun dan produktivitas pun terganggu."
    ))
    add_text(doc, (
        "Solusi konvensional seperti buku catatan atau aplikasi to-do list memerlukan "
        "upaya manual yang cukup besar untuk mengorganisir dan mengkategorikan setiap "
        "pikiran. Pengguna harus secara sadar memilah mana yang merupakan tugas, "
        "pengingat, target, ataupun sekadar catatan harian. Proses ini justru dapat "
        "menambah beban kognitif, terutama pada malam hari ketika energi mental "
        "pengguna sudah terkuras."
    ))
    add_text(doc, (
        "Kemajuan teknologi Kecerdasan Buatan (AI), khususnya Large Language Model (LLM), "
        "kini memungkinkan pemrosesan bahasa natural secara otomatis dengan tingkat "
        "akurasi yang tinggi. Model-model seperti DeepSeek v4-Flash mampu memahami "
        "konteks kalimat dan mengekstrak informasi terstruktur dari teks bebas. "
        "Kemampuan ini dapat dimanfaatkan untuk memproses catatan malam pengguna "
        "secara otomatis."
    ))
    add_text(doc, (
        "Night Dump hadir sebagai solusi inovatif yang mengintegrasikan teknologi AI "
        "dengan kebutuhan manajemen pikiran malam hari. Nama \"Night Dump\" sendiri "
        "terinspirasi dari istilah memory dump dalam dunia komputer, yaitu proses "
        "mengosongkan memori untuk keperluan analisis. Dalam konteks ini, pengguna "
        "dapat \"membuang\" (dump) semua pikiran mereka dalam format tulisan maupun "
        "suara, kemudian AI akan secara otomatis mengekstrak dan mengategorikannya "
        "menjadi tugas, pengingat, target, dan insight yang terorganisir."
    ))
    add_text(doc, (
        "Aplikasi ini dibangun menggunakan framework Flutter untuk mendukung "
        "pengembangan lintas platform yang efisien, dengan backend berupa database "
        "SQLite yang tersimpan secara lokal pada perangkat pengguna. Integrasi "
        "dengan API DeepSeek v4-Flash memastikan analisis AI yang akurat namun "
        "tetap hemat dalam penggunaan token, sehingga biaya operasional dapat "
        "ditekan seminimal mungkin."
    ))

    add_heading_2(doc, "1.2 Rumusan Masalah")
    add_text(doc, "Berdasarkan latar belakang di atas, rumusan masalah dalam proyek ini adalah:", indent=True)
    for item in [
        "Bagaimana merancang sistem input yang memudahkan pengguna mencatat pikiran malam secara bebas, baik melalui teks maupun suara?",
        "Bagaimana mengimplementasikan AI untuk mengekstrak dan mengategorikan item produktif (tugas, pengingat, target, insight) dari catatan bebas pengguna?",
        "Bagaimana membangun sistem manajemen tugas yang terintegrasi dengan hasil ekstraksi AI, termasuk fitur tenggat waktu (due date) otomatis?",
        "Bagaimana merancang antarmuka yang intuitif, estetis, dan nyaman digunakan pada malam hari dengan tema gelap?",
        "Bagaimana mengimplementasikan sistem autentikasi lokal yang aman untuk menjaga privasi data pengguna?",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "1.3 Tujuan Penelitian")
    add_text(doc, "Tujuan dari pengembangan aplikasi Night Dump adalah:", indent=True)
    for item in [
        "Membangun aplikasi mobile berbasis Flutter yang memungkinkan pengguna mencatat pikiran malam dengan mudah melalui teks atau input suara.",
        "Mengintegrasikan Large Language Model (DeepSeek v4-Flash) untuk menganalisis dan mengekstrak item produktif secara otomatis dari catatan pengguna.",
        "Mengimplementasikan sistem manajemen tugas dengan fitur tenggat waktu yang dapat diatur secara manual maupun ditetapkan otomatis oleh AI.",
        "Menyediakan dashboard analitik untuk melacak pola pikiran, distribusi item, dan tingkat produktivitas pengguna.",
        "Mengimplementasikan sistem autentikasi berbasis kredensial dengan password yang di-hash menggunakan algoritma SHA-256.",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "1.4 Manfaat Penelitian")
    add_heading_3(doc, "1.4.1 Manfaat Teoritis")
    for item in [
        "Memberikan kontribusi pada pengembangan aplikasi mobile yang mengintegrasikan AI untuk pemrosesan bahasa natural.",
        "Menjadi referensi implementasi praktis LLM API dalam konteks aplikasi produktivitas mobile.",
        "Menghasilkan kajian tentang efektivitas AI dalam membantu manajemen kognitif pengguna.",
    ]:
        add_bullet(doc, item)

    add_heading_3(doc, "1.4.2 Manfaat Praktis")
    for item in [
        "Membantu pengguna tidur lebih tenang dengan \"melepaskan\" beban pikiran secara terstruktur.",
        "Meningkatkan produktivitas pengguna melalui sistem manajemen tugas yang terorganisir otomatis.",
        "Memberikan insight tentang pola pikiran melalui analitik berbasis AI.",
        "Mempermudah pencatatan dengan fitur voice-to-text yang mendukung bahasa Indonesia.",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "1.5 Ruang Lingkup")
    add_text(doc, "Ruang lingkup pengembangan aplikasi Night Dump meliputi:", indent=True)
    for item in [
        "Platform target: Android (minimum SDK 21 / Android 5.0 Lollipop).",
        "Bahasa pemrograman: Dart dengan framework Flutter.",
        "Penyimpanan data: SQLite lokal pada perangkat (tidak memerlukan koneksi internet untuk data).",
        "Integrasi AI: DeepSeek v4-Flash API untuk analisis catatan (memerlukan koneksi internet).",
        "Autentikasi: Sistem login dan registrasi berbasis kredensial lokal.",
        "Input: Teks manual dan suara (Speech-to-Text bahasa Indonesia).",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BAB II TINJAUAN PUSTAKA
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "BAB II\nTINJAUAN PUSTAKA")

    add_heading_2(doc, "2.1 Mobile Programming")
    add_text(doc, (
        "Mobile programming adalah proses pengembangan perangkat lunak yang dirancang "
        "khusus untuk berjalan pada perangkat mobile seperti smartphone dan tablet. "
        "Aplikasi mobile memanfaatkan fitur-fitur unik perangkat seperti layar sentuh, "
        "GPS, kamera, mikrofon, dan sensor akselerometer untuk memberikan pengalaman "
        "pengguna yang imersif. Menurut Developers (2023), ekosistem mobile saat ini "
        "didominasi oleh dua platform utama: Android dan iOS."
    ))
    add_text(doc, (
        "Pendekatan pengembangan aplikasi mobile terbagi menjadi tiga kategori: "
        "native, hybrid, dan cross-platform. Pengembangan native menggunakan bahasa "
        "dan SDK resmi masing-masing platform (Kotlin/Java untuk Android, Swift/"
        "Objective-C untuk iOS), memberikan performa optimal namun memerlukan "
        "codebase terpisah. Pendekatan cross-platform seperti Flutter memungkinkan "
        "pengembangan satu codebase untuk multiple platform dengan performa mendekati native."
    ))

    add_heading_2(doc, "2.2 Flutter Framework")
    add_text(doc, (
        "Flutter adalah framework open-source yang dikembangkan oleh Google untuk "
        "membangun aplikasi multi-platform dari satu codebase tunggal. Flutter "
        "menggunakan bahasa pemrograman Dart yang dikompilasi langsung ke kode "
        "native ARM atau x86, menghasilkan performa yang setara dengan aplikasi native. "
        "Keunggulan utama Flutter meliputi: widget system yang konsisten di semua "
        "platform, hot reload untuk pengembangan yang cepat, dan ekosistem package "
        "yang kaya (pub.dev)."
    ))
    add_text(doc, (
        "Flutter menggunakan arsitektur berbasis widget di mana seluruh elemen UI "
        "adalah widget yang dapat disusun secara komposisi. State management dalam "
        "Flutter dapat dikelola menggunakan berbagai pendekatan: setState untuk state "
        "lokal, Provider, Riverpod, atau BLoC untuk state management yang lebih kompleks. "
        "Dalam Night Dump, state management menggunakan pendekatan setState yang "
        "sederhana namun efektif untuk kebutuhan aplikasi."
    ))

    add_heading_2(doc, "2.3 Large Language Model (LLM) dan Natural Language Processing")
    add_text(doc, (
        "Large Language Model (LLM) adalah jenis model kecerdasan buatan yang dilatih "
        "pada korpus teks besar untuk memahami dan menghasilkan bahasa natural. Model "
        "ini menggunakan arsitektur Transformer yang diperkenalkan oleh Vaswani et al. "
        "(2017), yang memungkinkan pemahaman konteks jangka panjang melalui mekanisme "
        "attention. LLM modern seperti GPT, Claude, dan DeepSeek telah menunjukkan "
        "kemampuan luar biasa dalam berbagai tugas NLP termasuk: klasifikasi teks, "
        "ekstraksi informasi, summarisasi, dan pemahaman instruksi kompleks."
    ))
    add_text(doc, (
        "DeepSeek v4-Flash, yang digunakan dalam Night Dump, adalah varian model "
        "DeepSeek yang dioptimalkan untuk kecepatan respons tinggi dengan biaya "
        "komputasi rendah. Model ini mendukung pemrosesan konteks hingga ratusan "
        "ribu token dan mampu memahami instruksi dalam berbagai bahasa termasuk "
        "Bahasa Indonesia. Kemampuan \"few-shot learning\" dari model ini memungkinkan "
        "aplikasi memberikan contoh format output yang diinginkan dalam prompt, "
        "sehingga konsistensi hasil ekstraksi dapat terjaga."
    ))

    add_heading_2(doc, "2.4 SQLite dan Manajemen Database Lokal")
    add_text(doc, (
        "SQLite adalah sistem manajemen database relasional (RDBMS) yang bersifat "
        "serverless, self-contained, dan zero-configuration. Database SQLite tersimpan "
        "dalam satu file pada perangkat, menjadikannya solusi ideal untuk penyimpanan "
        "data lokal pada aplikasi mobile. Menurut Hipp (2022), SQLite adalah database "
        "engine yang paling banyak digunakan di dunia, dengan lebih dari satu triliun "
        "deployment aktif."
    ))
    add_text(doc, (
        "Dalam Flutter, SQLite diakses melalui package sqflite yang menyediakan API "
        "asinkron untuk operasi database. Package ini mendukung transaksi, batch "
        "operation, dan foreign key constraints. Night Dump memanfaatkan fitur-fitur "
        "ini untuk memastikan integritas data antar tabel melalui mekanisme "
        "ON DELETE CASCADE pada relasi antara tabel notes dan note_items."
    ))

    add_heading_2(doc, "2.5 Speech-to-Text (STT)")
    add_text(doc, (
        "Speech-to-Text atau Automatic Speech Recognition (ASR) adalah teknologi "
        "yang mengkonversi sinyal audio ucapan manusia menjadi representasi teks. "
        "Teknologi ASR modern menggunakan model deep learning berbasis arsitektur "
        "sequence-to-sequence dengan mekanisme attention. Platform Android dan iOS "
        "menyediakan API ASR native yang memanfaatkan model bahasa lokal maupun "
        "cloud-based untuk akurasi tinggi."
    ))
    add_text(doc, (
        "Package speech_to_text pada Flutter menyediakan abstraksi lintas platform "
        "untuk mengakses layanan ASR native. Package ini mendukung berbagai bahasa "
        "termasuk Bahasa Indonesia (id_ID) dan menyediakan fitur partial results "
        "yang menampilkan teks secara real-time saat pengguna berbicara. Fitur ini "
        "meningkatkan pengalaman pengguna secara signifikan dibanding sistem "
        "yang hanya menampilkan hasil akhir."
    ))

    add_heading_2(doc, "2.6 Kriptografi: SHA-256 untuk Keamanan Password")
    add_text(doc, (
        "SHA-256 (Secure Hash Algorithm 256-bit) adalah fungsi hash kriptografis "
        "yang menghasilkan nilai hash 256-bit dari input data apapun. Fungsi hash "
        "bersifat satu arah (one-way), artinya tidak dapat dilakukan pembalikan "
        "dari nilai hash ke data aslinya. Dalam Night Dump, SHA-256 digunakan untuk "
        "meng-hash password pengguna sebelum disimpan ke database, sehingga bahkan "
        "jika database lokal diakses oleh pihak yang tidak berwenang, password "
        "asli pengguna tetap terlindungi."
    ))

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BAB III ANALISIS DAN PERANCANGAN
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "BAB III\nANALISIS DAN PERANCANGAN SISTEM")

    add_heading_2(doc, "3.1 Analisis Kebutuhan Sistem")
    add_heading_3(doc, "3.1.1 Kebutuhan Fungsional")
    add_text(doc, "Sistem Night Dump harus mampu:", indent=True)
    for item in [
        "Autentikasi Pengguna: Registrasi akun baru, login dengan email dan password, serta logout yang menghapus sesi aktif.",
        "Input Catatan: Menerima input teks bebas dan input suara (speech-to-text) dari pengguna.",
        "Analisis AI: Mengirimkan catatan ke AI DeepSeek dan mendapatkan respons berupa ringkasan dan daftar item terstruktur.",
        "Ekstraksi Item: Mengekstrak item bertype todo, reminder, target, dan insight beserta tenggat waktu (due_date) jika disebutkan.",
        "Penyimpanan Data: Menyimpan catatan dan item hasil ekstraksi ke database SQLite lokal.",
        "Manajemen Tugas: Menampilkan, menandai selesai, mengatur tenggat waktu, dan menghapus item per kategori.",
        "Histori Catatan: Menampilkan riwayat catatan dengan item-item yang telah diekstrak AI.",
        "Insight Analitik: Menampilkan statistik catatan, distribusi item, tingkat penyelesaian tugas, dan insight terbaru.",
        "Mood Tracking: Mencatat mood malam pengguna yang digunakan sebagai konteks tambahan untuk AI.",
    ]:
        add_bullet(doc, item)

    add_heading_3(doc, "3.1.2 Kebutuhan Non-Fungsional")
    for item in [
        "Keamanan: Password pengguna di-hash dengan SHA-256 sebelum disimpan; data tersimpan lokal tanpa transmisi ke server eksternal.",
        "Performa: Respons UI tidak lebih dari 100ms untuk operasi lokal; timeout AI ditetapkan maksimal 45 detik.",
        "Ketersediaan Offline: Seluruh fitur manajemen tugas dan histori tersedia tanpa koneksi internet; hanya fitur AI yang memerlukan internet.",
        "Usability: Antarmuka menggunakan tema gelap (dark mode) yang nyaman untuk penggunaan malam hari.",
        "Kompatibilitas: Mendukung Android versi 5.0 (API 21) ke atas.",
        "Skalabilitas: Struktur database dirancang modular untuk mendukung penambahan fitur di masa mendatang.",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "3.2 Arsitektur Sistem")
    add_text(doc, (
        "Night Dump menggunakan arsitektur berlapis (layered architecture) yang terdiri dari tiga lapisan utama:"
    ))
    for item in [
        "Presentation Layer (UI): Halaman-halaman Flutter yang menampilkan antarmuka pengguna, terdiri dari Home, Notes, History, Insights, Tasks, Login, dan Register.",
        "Business Logic Layer (Services): Kelas-kelas service yang menangani logika bisnis, yaitu AuthService (autentikasi), AiService (integrasi AI), dan DatabaseHelper (manajemen database).",
        "Data Layer (Database): SQLite database dengan empat tabel utama: users, notes, note_items, dan settings.",
    ]:
        add_bullet(doc, item)
    add_text(doc, (
        "Komunikasi antar lapisan bersifat satu arah ke bawah: Presentation Layer memanggil "
        "Business Logic Layer, yang kemudian berinteraksi dengan Data Layer. Lapisan presentasi "
        "tidak pernah mengakses database secara langsung, memastikan separation of concerns "
        "yang baik. Integrasi dengan DeepSeek API dilakukan melalui AiService "
        "menggunakan protokol HTTP dengan format JSON yang kompatibel dengan OpenAI API."
    ))

    add_heading_2(doc, "3.3 Rancangan Database")
    add_text(doc, "Database Night Dump terdiri dari empat tabel dengan relasi sebagai berikut:")

    # Tabel users
    add_heading_3(doc, "3.3.1 Tabel users")
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['Kolom', 'Tipe Data', 'Keterangan']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    data = [
        ('id', 'INTEGER PK AUTOINCREMENT', 'Primary key unik'),
        ('name', 'TEXT NOT NULL', 'Nama lengkap pengguna'),
        ('email', 'TEXT UNIQUE NOT NULL', 'Alamat email (unik)'),
        ('password_hash', 'TEXT NOT NULL', 'Hash SHA-256 dari password'),
        ('created_at', 'TEXT NOT NULL', 'Waktu registrasi (ISO 8601)'),
    ]
    for i, (col, dtype, desc) in enumerate(data):
        row = table.rows[i+1]
        for j, val in enumerate([col, dtype, desc]):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            set_font(run, 11)
    add_spacer(doc)

    # Tabel notes
    add_heading_3(doc, "3.3.2 Tabel notes")
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    data2 = [
        ('id', 'INTEGER PK AUTOINCREMENT', 'Primary key catatan'),
        ('user_email', 'TEXT NOT NULL', 'Email pengguna pemilik catatan'),
        ('raw_text', 'TEXT NOT NULL', 'Teks asli catatan pengguna'),
        ('summary', 'TEXT DEFAULT ""', 'Ringkasan hasil AI'),
        ('created_at', 'TEXT NOT NULL', 'Waktu catatan dibuat (ISO 8601)'),
    ]
    for i, (col, dtype, desc) in enumerate(data2):
        row = table2.rows[i+1]
        for j, val in enumerate([col, dtype, desc]):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            set_font(run, 11)
    add_spacer(doc)

    # Tabel note_items
    add_heading_3(doc, "3.3.3 Tabel note_items")
    table3 = doc.add_table(rows=7, cols=3)
    table3.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table3.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    data3 = [
        ('id', 'INTEGER PK AUTOINCREMENT', 'Primary key item'),
        ('note_id', 'INTEGER FK → notes.id', 'Relasi ke tabel notes'),
        ('type', 'TEXT NOT NULL', 'Kategori: todo/reminder/target/insight'),
        ('content', 'TEXT NOT NULL', 'Isi item hasil ekstraksi AI'),
        ('done', 'INTEGER DEFAULT 0', 'Status penyelesaian (0=aktif, 1=selesai)'),
        ('due_date', 'TEXT (nullable)', 'Tenggat waktu format "YYYY-MM-DD HH:mm"'),
    ]
    for i, (col, dtype, desc) in enumerate(data3):
        row = table3.rows[i+1]
        for j, val in enumerate([col, dtype, desc]):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            set_font(run, 11)
    add_spacer(doc)

    # Tabel settings
    add_heading_3(doc, "3.3.4 Tabel settings")
    table4 = doc.add_table(rows=3, cols=3)
    table4.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table4.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    data4 = [
        ('key', 'TEXT PRIMARY KEY', 'Kunci pengaturan (misal: current_user, last_mood)'),
        ('value', 'TEXT (nullable)', 'Nilai pengaturan dalam format teks'),
    ]
    for i, (col, dtype, desc) in enumerate(data4):
        row = table4.rows[i+1]
        for j, val in enumerate([col, dtype, desc]):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            set_font(run, 11)

    add_text(doc, (
        "Tabel settings berfungsi sebagai key-value store yang fleksibel. "
        "Kunci 'current_user' menyimpan email pengguna yang sedang login untuk "
        "manajemen sesi, sedangkan kunci 'last_mood' menyimpan mood terakhir yang "
        "dipilih pengguna sebagai konteks tambahan untuk prompt AI."
    ))

    add_heading_2(doc, "3.4 Alur Kerja Sistem (Workflow)")

    add_heading_3(doc, "3.4.1 Alur Autentikasi")
    add_text(doc, "Alur autentikasi terdiri dari dua proses utama:", indent=True)
    add_text(doc, "Registrasi:", bold=False, indent=True)
    for step in [
        "Pengguna mengisi formulir nama, email, dan password.",
        "Sistem memvalidasi format email dan panjang password (minimal 6 karakter).",
        "Sistem memeriksa apakah email sudah terdaftar dalam tabel users.",
        "Password di-hash menggunakan SHA-256.",
        "Data pengguna disimpan ke tabel users.",
        "Sesi aktif dibuat dengan menyimpan email ke tabel settings (key: current_user).",
        "Pengguna diarahkan ke halaman Home.",
    ]:
        add_bullet(doc, step)

    add_text(doc, "Login:", bold=False, indent=True)
    for step in [
        "Pengguna memasukkan email dan password.",
        "Sistem mencari pengguna berdasarkan email di tabel users.",
        "Password yang dimasukkan di-hash dan dibandingkan dengan hash yang tersimpan.",
        "Jika cocok, sesi aktif dibuat dan pengguna diarahkan ke Home.",
        "Jika tidak cocok, pesan error ditampilkan.",
    ]:
        add_bullet(doc, step)

    add_heading_3(doc, "3.4.2 Alur Input dan Analisis Catatan")
    for step in [
        "Pengguna membuka halaman Catatan Baru dari navigasi bawah (tombol +).",
        "Pengguna memilih mood malam ini di halaman Home (opsional; disimpan ke tabel settings).",
        "Pengguna mengetik catatan bebas di area teks, atau menekan tombol mikrofon untuk input suara.",
        "Speech-to-Text (speech_to_text) mengkonversi ucapan menjadi teks secara real-time.",
        "Pengguna menekan tombol LEPASKAN untuk mengirimkan catatan.",
        "Sistem membaca mood terakhir dari tabel settings.",
        "Sistem memanggil AiService.analyzeNote() dengan teks catatan dan mood sebagai parameter.",
        "AiService mengirimkan HTTP POST ke DeepSeek API dengan sistem prompt yang menyertakan tanggal hari ini, referensi waktu relatif (besok, lusa), dan mood pengguna.",
        "DeepSeek memproses catatan dan mengembalikan respons JSON berisi summary dan daftar items.",
        "AiService mem-parsing respons JSON menjadi objek AiResult berisi list NoteItem.",
        "Jika parsing berhasil, sistem menyimpan catatan ke tabel notes dan setiap item ke tabel note_items.",
        "Pengguna diarahkan ke halaman History.",
        "Jika AI gagal, ditampilkan bottom sheet dengan pilihan: Simpan Saja atau Coba Lagi.",
    ]:
        add_bullet(doc, step)

    add_heading_3(doc, "3.4.3 Alur Manajemen Tugas")
    for step in [
        "Pengguna membuka halaman Tugas melalui ikon checklist di navigasi bawah.",
        "Halaman menampilkan tiga tab: TUGAS, PENGINGAT, dan TARGET.",
        "Setiap tab memuat item-item yang relevan dari tabel note_items, diurutkan berdasarkan: item aktif terlebih dahulu, kemudian yang memiliki due_date terdekat di atas.",
        "Pengguna dapat menandai item sebagai selesai dengan menekan ikon centang (toggle done).",
        "Pengguna dapat mengatur atau mengubah tenggat waktu dengan menekan chip due_date (membuka DatePicker dan TimePicker).",
        "Pengguna dapat menghapus item dengan gesture swipe ke kiri (Dismissible).",
        "Item yang telah selesai ditampilkan di bagian bawah dengan teks tercoret.",
    ]:
        add_bullet(doc, step)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BAB IV IMPLEMENTASI
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "BAB IV\nIMPLEMENTASI SISTEM")

    add_heading_2(doc, "4.1 Teknologi dan Tools yang Digunakan")

    # Tabel teknologi
    table_tech = doc.add_table(rows=10, cols=3)
    table_tech.style = 'Table Grid'
    for i, h in enumerate(['Teknologi/Library', 'Versi', 'Fungsi']):
        cell = table_tech.rows[0].cells[i]
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        set_font(r, 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tech_data = [
        ('Flutter', '3.x', 'Framework utama pengembangan UI cross-platform'),
        ('Dart', '3.x', 'Bahasa pemrograman'),
        ('sqflite', '^2.4.1', 'Akses database SQLite lokal'),
        ('path', '^1.9.1', 'Manajemen path file database'),
        ('http', '^1.2.2', 'HTTP client untuk panggilan DeepSeek API'),
        ('crypto', '^3.0.3', 'Hashing SHA-256 untuk keamanan password'),
        ('speech_to_text', '^7.0.0', 'Input suara (speech recognition)'),
        ('DeepSeek API', 'v4-Flash', 'Model AI untuk analisis dan ekstraksi teks'),
        ('SQLite', '3.x', 'Database relasional lokal pada perangkat'),
    ]
    for i, (tech, ver, func) in enumerate(tech_data):
        row = table_tech.rows[i+1]
        for j, val in enumerate([tech, ver, func]):
            row.cells[j].text = ''
            run = row.cells[j].paragraphs[0].add_run(val)
            set_font(run, 11)
    add_spacer(doc)

    add_heading_2(doc, "4.2 Struktur Navigasi Aplikasi")
    add_text(doc, (
        "Navigasi Night Dump menggunakan sistem named routes yang terdefinisi di main.dart. "
        "Terdapat tujuh rute utama: /login, /register, /home, /notes, /history, /insights, dan /tasks. "
        "Bar navigasi bawah (AppNavBar) yang digunakan bersama di semua halaman utama "
        "memanfaatkan named routes untuk menghindari circular import antar halaman."
    ))
    add_text(doc, "Deskripsi setiap halaman utama:", indent=True)
    for item in [
        "Home (/home): Dashboard utama yang menampilkan greeting berbasis waktu, statistik harian (catatan, tugas aktif, pengingat), pemilih mood, dan tiga catatan terbaru.",
        "Catatan (/notes): Halaman input catatan baru dengan dukungan teks dan suara, loading state dengan copy poetik, dan penanganan error AI.",
        "Histori (/history): Daftar semua catatan pengguna dengan tampilan item-item hasil ekstraksi AI yang dikategorikan per tipe.",
        "Insight (/insights): Dashboard analitik dengan statistik catatan, distribusi item (bar chart dengan LinearProgressIndicator), ring chart penyelesaian tugas, dan daftar insight terbaru.",
        "Tugas (/tasks): Halaman manajemen tugas dengan tiga tab (Tugas, Pengingat, Target), fitur centang selesai, pengaturan due_date, dan swipe-to-delete.",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "4.3 Implementasi Sistem AI")

    add_heading_3(doc, "4.3.1 Desain Prompt (Prompt Engineering)")
    add_text(doc, (
        "Efektivitas ekstraksi AI sangat bergantung pada desain prompt yang digunakan. "
        "Night Dump menggunakan teknik few-shot prompting di mana contoh konkret "
        "disertakan dalam sistem prompt untuk memandu model menghasilkan output "
        "dalam format yang tepat. Sistem prompt dibangun secara dinamis dengan "
        "menyertakan:"
    ))
    for item in [
        "Tanggal hari ini dalam format YYYY-MM-DD beserta nama hari (misal: Minggu, 2026-06-01).",
        "Referensi waktu relatif yang sudah dihitung: tanggal besok dan lusa.",
        "Konvensi waktu dalam sehari: pagi≈07:00, siang≈12:00, sore≈17:00, malam≈20:00.",
        "Mood pengguna (jika tersedia) sebagai konteks emosional.",
        "Contoh output JSON konkret dengan format yang diharapkan.",
    ]:
        add_bullet(doc, item)

    add_text(doc, (
        "Format output yang diminta adalah JSON dengan struktur: "
        '{"summary": "ringkasan catatan", "items": [{"type": "todo|reminder|target|insight", '
        '"content": "isi item", "due_date": "YYYY-MM-DD HH:mm atau null"}]}. '
        "Penggunaan contoh konkret (few-shot) terbukti lebih reliabel dibanding "
        "mendeskripsikan format secara abstrak, karena model dapat mengikuti "
        "pola yang sudah tersedia secara langsung."
    ))

    add_heading_3(doc, "4.3.2 Penanganan Respons AI")
    add_text(doc, "Respons dari DeepSeek API diproses melalui beberapa tahap validasi:", indent=True)
    for step in [
        "Stripping markdown: Model terkadang membungkus JSON dalam markdown code block (```json ... ```). Regex digunakan untuk menghapus pembungkus ini.",
        "Ekstraksi JSON: Jika terdapat teks di luar JSON, regex diterapkan untuk mengekstrak blok { ... } terluar.",
        "Parsing JSON: jsonDecode() digunakan untuk mengurai string JSON menjadi struktur data Dart.",
        "Validasi tipe: Setiap item divalidasi untuk memastikan field type dan content ada dan tidak kosong.",
        "Normalisasi due_date: Nilai null (JSON null), string 'null', dan string kosong dikonversi ke null Dart yang sesungguhnya.",
        "Fallback: Jika parsing gagal pada tahap apapun, fungsi _parse() mengembalikan null dan UI menampilkan dialog error dengan opsi retry atau save-without-AI.",
    ]:
        add_bullet(doc, step)

    add_heading_3(doc, "4.3.3 Parameter API")
    add_text(doc, "Konfigurasi request ke DeepSeek API adalah sebagai berikut:", indent=True)
    for item in [
        "Model: deepseek-v4-flash (model berbiaya rendah dengan respons cepat)",
        "max_tokens: 2000 (cukup untuk mengekstrak 30-40 item dari catatan panjang)",
        "temperature: 0.3 (mendekati deterministik untuk konsistensi format output)",
        "stream: false (respons diterima sekaligus, bukan streaming)",
        "Timeout: 45 detik (mengakomodasi input suara panjang)",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "4.4 Implementasi Fitur Voice Input")
    add_text(doc, (
        "Fitur voice input menggunakan package speech_to_text v7.0.0 yang memanfaatkan "
        "layanan ASR native dari sistem operasi Android dan iOS. Implementasi dirancang "
        "untuk mendukung sesi rekaman panjang (hingga 5 menit per sesi) dengan "
        "beberapa mekanisme:"
    ))
    for item in [
        "Partial Results: Teks ditampilkan secara real-time saat pengguna berbicara, memberikan feedback visual instan.",
        "Auto-Resume: Setiap sesi rekaman baru akan menyimpan teks yang sudah ada sebagai 'base text', sehingga rekaman berikutnya ditambahkan (append) ke bawah.",
        "Auto-Stop: Speech recognition berhenti otomatis setelah 3 detik keheningan (pauseFor) atau setelah 5 menit (listenFor).",
        "Error Handling: Kesalahan mikrofon ditangani dengan mengembalikan state _isListening ke false dan menampilkan notifikasi.",
        "Locale: ID Indonesia (id_ID) dikonfigurasi sebagai locale utama untuk akurasi pengenalan Bahasa Indonesia.",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "4.5 Antarmuka Pengguna (UI/UX)")
    add_text(doc, (
        "Night Dump menggunakan desain visual dark-themed dengan palet warna utama: "
        "latar belakang ungu gelap (#13132B), aksen lavender (#D1B3FF), dan "
        "elemen kartu berwarna (#1A1A3A). Filosofi desain mengutamakan kenyamanan "
        "visual pada malam hari dengan kontras yang seimbang dan mengurangi emisi "
        "cahaya biru yang berlebihan."
    ))
    add_text(doc, "Prinsip-prinsip desain yang diterapkan:", indent=True)
    for item in [
        "Floating Navigation Bar: Bar navigasi bawah menggunakan desain pill shape (rounded rectangle) yang mengambang di atas konten, dengan efek glowing pada tombol aktif.",
        "Glassmorphism Cards: Kartu-kartu konten menggunakan efek semi-transparan dengan border tipis untuk kedalaman visual.",
        "Animated Transitions: AnimatedSwitcher digunakan untuk transisi konten yang halus dengan efek fade dan slide.",
        "Color Coding: Setiap tipe item memiliki warna identitas: ungu (tugas), oranye (pengingat), hijau (target), biru (insight).",
        "Loading Copywriting: Pesan loading menggunakan bahasa poetik yang tidak menyebutkan AI, menciptakan pengalaman yang lebih personal.",
        "Responsive Layout: Layout menggunakan Expanded, Flexible, dan SingleChildScrollView untuk adaptasi berbagai ukuran layar.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # BAB V PENUTUP
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "BAB V\nPENUTUP")

    add_heading_2(doc, "5.1 Kesimpulan")
    add_text(doc, (
        "Night Dump merupakan aplikasi mobile inovatif yang berhasil mengintegrasikan "
        "teknologi Flutter, SQLite, dan Large Language Model DeepSeek untuk "
        "menciptakan solusi manajemen pikiran malam yang komprehensif. Beberapa "
        "pencapaian utama dari pengembangan aplikasi ini:"
    ))
    for item in [
        "Berhasil mengimplementasikan alur end-to-end dari input bebas (teks/suara) hingga ekstraksi terstruktur (todo, reminder, target, insight) menggunakan AI.",
        "Sistem prompt few-shot yang efisien memungkinkan AI memahami konteks waktu relatif (besok, lusa) dan mengkonversinya ke tenggat waktu konkret.",
        "Arsitektur database yang modular dengan empat tabel (users, notes, note_items, settings) mendukung integritas data dan skalabilitas.",
        "Fitur voice input dengan partial results real-time meningkatkan aksesibilitas dan kemudahan penggunaan.",
        "Sistem autentikasi lokal dengan hashing SHA-256 menjamin keamanan data pengguna tanpa ketergantungan server eksternal.",
        "Antarmuka dark-themed yang estetis dan nyaman untuk penggunaan malam hari.",
    ]:
        add_bullet(doc, item)

    add_heading_2(doc, "5.2 Saran Pengembangan")
    add_text(doc, "Beberapa area pengembangan yang dapat dilakukan untuk meningkatkan aplikasi:", indent=True)
    for item in [
        "Integrasi Notifikasi: Implementasi push notification untuk pengingat (reminder) dan tugas yang mendekati tenggat waktu.",
        "Sinkronisasi Cloud: Penambahan fitur backup dan sinkronisasi data ke cloud storage untuk keamanan data.",
        "Analitik Lanjutan: Implementasi grafik tren mood mingguan dan word cloud dari tema-tema catatan menggunakan data historis.",
        "Multi-bahasa: Dukungan bahasa tambahan (Inggris, dll.) untuk jangkauan pengguna yang lebih luas.",
        "Widget Home Screen: Pembuatan widget Android untuk akses cepat ke catatan terbaru dari layar beranda.",
        "Offline AI: Eksplorasi penggunaan model AI yang dapat berjalan on-device untuk mengurangi ketergantungan koneksi internet.",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # DAFTAR PUSTAKA
    # ═══════════════════════════════════════════════════════════════════════════
    add_heading_1(doc, "DAFTAR PUSTAKA")

    refs = [
        "DeepSeek. (2024). DeepSeek API Documentation. DeepSeek Inc. Diakses dari https://platform.deepseek.com/docs",
        "Flutter Team. (2023). Flutter Documentation: Build apps for any screen. Google LLC. Diakses dari https://flutter.dev/docs",
        "Google Developers. (2023). Android Developer Guide. Google LLC. Diakses dari https://developer.android.com/guide",
        "Hipp, D. R. (2022). SQLite Documentation. SQLite Consortium. Diakses dari https://www.sqlite.org/docs.html",
        "Richards, J. (2022). speech_to_text: A Flutter plugin for speech recognition. pub.dev. Diakses dari https://pub.dev/packages/speech_to_text",
        "Vaswani, A., Shazeer, N., Parmar, N., et al. (2017). Attention is All You Need. Advances in Neural Information Processing Systems, 30.",
        "Brown, T., Mann, B., Ryder, N., et al. (2020). Language Models are Few-Shot Learners. Advances in Neural Information Processing Systems, 33.",
        "sqflite Contributors. (2023). sqflite: SQLite plugin for Flutter. pub.dev. Diakses dari https://pub.dev/packages/sqflite",
        "National Institute of Standards and Technology. (2015). FIPS PUB 180-4: Secure Hash Standard (SHS). U.S. Department of Commerce.",
        "Dart Team. (2023). Dart Documentation. Google LLC. Diakses dari https://dart.dev/guides",
    ]
    for i, ref in enumerate(refs, 1):
        para = doc.add_paragraph()
        run = para.add_run(ref)
        set_font(run, 12)
        pf = para.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        pf.first_line_indent = Cm(-1.25)
        pf.left_indent = Cm(1.25)

    # ── Simpan file ──
    filename = 'proposal_night_dump.docx'
    doc.save(filename)
    print(f"✅ File berhasil dibuat: {filename}")
    print(f"   Lokasi: {__import__('os').path.abspath(filename)}")


if __name__ == '__main__':
    create_proposal()
